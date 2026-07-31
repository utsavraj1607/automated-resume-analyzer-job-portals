import os
import re
import json
import unicodedata
from datetime import datetime

import spacy
from pdfminer.high_level import extract_text as pdf_extract_text
import docx


# --------------------------------------------------------------------------
# Constants / Regex patterns
# --------------------------------------------------------------------------

EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")

# Matches Indian & International phone numbers with optional country code,
# spaces, dashes, parentheses.
PHONE_RE = re.compile(
    r"(?:(?:\+|00)\d{1,3}[\s.-]?)?"          # optional country code
    r"(?:\(\d{2,4}\)[\s.-]?)?"               # optional area code in parens
    r"(?:\d{3,5}[\s.-]?){2,4}"               # groups of digits
)

LINKEDIN_RE = re.compile(r"(https?://)?(www\.)?linkedin\.com/[A-Za-z0-9\-_/%]+", re.I)
GITHUB_RE = re.compile(r"(https?://)?(www\.)?github\.com/[A-Za-z0-9\-_/%]+", re.I)
GENERIC_URL_RE = re.compile(r"(https?://[^\s,;]+)", re.I)

DATE_RANGE_RE = re.compile(
    r"(?P<start>(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{4})"
    r"\s*(?:-|–|—|to)\s*"
    r"(?P<end>(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{4}|[Pp]resent|[Cc]urrent)",
    re.I,
)

DEGREE_KEYWORDS = [
    "Bachelor", "Master", "B.Tech", "M.Tech", "B.E", "M.E", "B.Sc", "M.Sc",
    "BSc", "MSc", "MBA", "PhD", "Ph.D", "B.A", "M.A", "BA", "MA", "B.Com",
    "M.Com", "Diploma", "Associate Degree", "Doctorate", "B.C.A", "M.C.A",
    "BCA", "MCA",
]

SECTION_HEADERS = {
    "summary": ["summary", "objective", "professional summary", "profile", "about me"],
    "experience": [
        "experience", "work experience", "professional experience",
        "work history", "employment history", "employment",
    ],
    "education": ["education", "academic background", "academic qualifications", "qualifications"],
    "skills": ["skills", "technical skills", "core competencies", "skill set", "key skills"],
    "projects": ["projects", "academic projects", "personal projects", "project experience"],
    "certifications": ["certifications", "certificates", "licenses"],
    "contact": ["contact", "contact information", "personal details"],
}


class ResumeParser:
    """
    Ingests a resume file (PDF or DOCX), extracts and cleans the raw text,
    segments it into logical sections, extracts structured entities, and
    returns a standardized JSON-serializable dictionary.
    """

    def __init__(self, skills_db_path: str = None, spacy_model: str = "en_core_web_sm"):
        if skills_db_path is None:
            skills_db_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "skills_db.json")

        with open(skills_db_path, "r", encoding="utf-8") as f:
            self.skills_db = json.load(f)

        # Flatten skill ontology into a single lookup list (skill -> category)
        self.skill_lookup = {}
        for category, skills in self.skills_db.items():
            for skill in skills:
                self.skill_lookup[skill.lower()] = {"canonical": skill, "category": category}

        # Add common aliases so variations map back to a canonical skill name
        self.aliases = {
            "reactjs": "React", "react js": "React", "react.js": "React",
            "nodejs": "Node.js", "node js": "Node.js",
            "vuejs": "Vue.js", "vue js": "Vue.js",
            "nextjs": "Next.js", "next js": "Next.js",
            "ml": "Machine Learning", "dl": "Deep Learning",
            "nlp": "NLP", "cv": "Computer Vision",
            "aws": "AWS", "gcp": "GCP",
            "js": "JavaScript", "ts": "TypeScript",
            "oop": "OOP", "dsa": "DSA", "sql": "SQL",
        }
        for alias, canonical in self.aliases.items():
            canonical_lower = canonical.lower()
            if canonical_lower in self.skill_lookup:
                self.skill_lookup[alias] = self.skill_lookup[canonical_lower]

        try:
            self.nlp = spacy.load(spacy_model)
        except OSError as exc:
            raise OSError(
                f"spaCy model '{spacy_model}' is not installed. "
                f"Run: python -m spacy download {spacy_model}"
            ) from exc

    # ----------------------------------------------------------------
    # 1. Document ingestion & text extraction
    # ----------------------------------------------------------------

    def _extract_text_from_pdf(self, file_path: str) -> str:
        return pdf_extract_text(file_path)

    def _extract_text_from_docx(self, file_path: str) -> str:
        document = docx.Document(file_path)
        parts = []
        for para in document.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Also capture text inside tables (common in resume templates)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)

    def extract_raw_text(self, file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return self._extract_text_from_pdf(file_path)
        elif ext in (".docx", ".doc"):
            return self._extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}. Only .pdf and .docx are supported.")

    def clean_text(self, text: str) -> str:
        """Normalize whitespace, strip non-ascii noise, standardize bullets."""
        # Normalize unicode (e.g. curly quotes, accented chars) to closest ascii
        text = unicodedata.normalize("NFKD", text)

        # Standardize bullet characters to '-'
        text = re.sub(r"[•●▪◦‣∙·]", "-", text)

        # Remove page-number-only lines e.g. "Page 1 of 2" / lone numbers
        text = re.sub(r"(?im)^\s*page\s+\d+(\s+of\s+\d+)?\s*$", "", text)

        # Collapse multiple blank lines
        text = re.sub(r"\n\s*\n+", "\n\n", text)

        # Collapse repeated spaces/tabs (but keep newlines)
        text = re.sub(r"[ \t]+", " ", text)

        # Strip trailing/leading whitespace on each line
        lines = [line.strip() for line in text.split("\n")]
        text = "\n".join(lines)

        return text.strip()

    # ----------------------------------------------------------------
    # 2. Section segmentation
    # ----------------------------------------------------------------

    def _match_header(self, line: str):
        """Return section key if `line` looks like a section heading, else None."""
        candidate = line.strip().lower().strip(":").strip()
        if not candidate or len(candidate.split()) > 5:
            return None
        for section, keywords in SECTION_HEADERS.items():
            for kw in keywords:
                if candidate == kw or candidate.startswith(kw):
                    return section
        return None

    def segment_sections(self, text: str) -> dict:
        """
        Splits cleaned resume text into logical sections using keyword-based
        heuristics. Any text before the first recognized header is treated
        as the 'header' block (typically name + contact info).
        """
        lines = text.split("\n")
        sections = {"header": []}
        current_section = "header"

        for line in lines:
            header_match = self._match_header(line)
            if header_match:
                current_section = header_match
                sections.setdefault(current_section, [])
                continue
            sections.setdefault(current_section, []).append(line)

        # Join back into text blocks, drop empties
        return {k: "\n".join(v).strip() for k, v in sections.items() if "\n".join(v).strip()}

    # ----------------------------------------------------------------
    # 3. Information extraction
    # ----------------------------------------------------------------

    def extract_contact_info(self, full_text: str) -> dict:
        email_match = EMAIL_RE.search(full_text)
        linkedin_match = LINKEDIN_RE.search(full_text)
        github_match = GITHUB_RE.search(full_text)

        # Phone: search candidates, keep the one with 7-15 digits
        phone = None
        for match in PHONE_RE.finditer(full_text):
            candidate = match.group().strip()
            digits = re.sub(r"\D", "", candidate)
            if 7 <= len(digits) <= 15:
                phone = self._normalize_phone(candidate)
                break

        portfolio = None
        for url_match in GENERIC_URL_RE.finditer(full_text):
            url = url_match.group()
            if "linkedin.com" not in url and "github.com" not in url:
                portfolio = url
                break

        return {
            "email": email_match.group() if email_match else None,
            "phone": phone,
            "linkedin": linkedin_match.group() if linkedin_match else None,
            "github": github_match.group() if github_match else None,
            "portfolio": portfolio,
        }

    def _normalize_phone(self, raw_phone: str) -> str:
        """Normalize a phone number to a standard '+<countrycode> <number>' format."""
        digits = re.sub(r"\D", "", raw_phone)
        if raw_phone.strip().startswith("+"):
            return "+" + digits
        if digits.startswith("00"):
            return "+" + digits[2:]
        if len(digits) == 10:
            # Assume domestic number, default to India (+91) per project context;
            # falls back gracefully if this assumption doesn't hold.
            return "+91 " + digits
        return "+" + digits

    def extract_name(self, header_text: str):
        """Uses spaCy NER on the header block (top of resume) to find the candidate name."""
        doc = self.nlp(header_text[:300])  # name is virtually always near the top
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                return ent.text.strip()
        # Fallback: assume the first non-empty line is the name if short
        first_line = header_text.strip().split("\n")[0].strip()
        if 0 < len(first_line.split()) <= 4 and not EMAIL_RE.search(first_line):
            return first_line
        return None

    def extract_education(self, education_text: str) -> list:
        """
        Parses the education section into a list of entries, each with
        degree, institution, and dates (where detectable).
        """
        if not education_text:
            return []

        entries = []
        blocks = [b.strip() for b in re.split(r"\n{1,}", education_text) if b.strip()]
        doc = self.nlp(education_text)

        # Map organizations detected by spaCy for institution lookup
        orgs_in_text = [ent.text for ent in doc.ents if ent.label_ == "ORG"]

        current_entry = None
        for line in blocks:
            degree_found = next((d for d in DEGREE_KEYWORDS if d.lower() in line.lower()), None)
            date_match = DATE_RANGE_RE.search(line)

            if degree_found or date_match:
                if current_entry:
                    entries.append(current_entry)
                current_entry = {
                    "degree": degree_found,
                    "institution": None,
                    "start_date": date_match.group("start") if date_match else None,
                    "end_date": date_match.group("end") if date_match else None,
                    "raw_line": line,
                }
                # Try to find an institution name in this same line
                for org in orgs_in_text:
                    if org in line:
                        current_entry["institution"] = org
                        break
            elif current_entry is not None:
                # Likely a continuation line (e.g. institution name, GPA)
                if current_entry["institution"] is None:
                    for org in orgs_in_text:
                        if org in line:
                            current_entry["institution"] = org
                            break
                    if current_entry["institution"] is None:
                        current_entry["institution"] = line
            else:
                current_entry = {
                    "degree": None,
                    "institution": line,
                    "start_date": None,
                    "end_date": None,
                    "raw_line": line,
                }

        if current_entry:
            entries.append(current_entry)

        return entries

    def extract_experience(self, experience_text: str) -> list:
        """
        Parses the experience section into a list of job entries, each with
        company, title, dates, and description bullet points.
        """
        if not experience_text:
            return []

        lines = [l for l in experience_text.split("\n") if l.strip()]
        doc = self.nlp(experience_text)
        orgs_in_text = [ent.text for ent in doc.ents if ent.label_ == "ORG"]

        entries = []
        current_entry = None

        for line in lines:
            date_match = DATE_RANGE_RE.search(line)
            is_bullet = line.strip().startswith("-")

            if date_match and not is_bullet:
                # New job entry starts here
                if current_entry:
                    entries.append(current_entry)

                # Strip the date substring to isolate company/title text
                remainder = DATE_RANGE_RE.sub("", line).strip(" -,|")
                company = next((org for org in orgs_in_text if org in line), None)
                title = remainder.replace(company, "").strip(" -,|") if company else remainder

                current_entry = {
                    "company": company,
                    "job_title": title if title else None,
                    "start_date": date_match.group("start"),
                    "end_date": date_match.group("end"),
                    "description": [],
                }
            elif is_bullet and current_entry is not None:
                current_entry["description"].append(line.lstrip("- ").strip())
            elif current_entry is not None and current_entry.get("job_title") is None:
                current_entry["job_title"] = line.strip()
            elif current_entry is not None:
                current_entry["description"].append(line.strip())
            else:
                # No date/header seen yet; start a provisional entry
                current_entry = {
                    "company": None,
                    "job_title": line.strip(),
                    "start_date": None,
                    "end_date": None,
                    "description": [],
                }

        if current_entry:
            entries.append(current_entry)

        return entries

    def extract_skills(self, full_text: str) -> list:
        """
        Matches resume text against the skill ontology (skills_db.json),
        handling common phrasing variations via the alias table.
        """
        found = {}
        text_lower = full_text.lower()

        for key, meta in self.skill_lookup.items():
            # Word-boundary match to avoid partial substring false positives,
            # e.g. "R" matching inside "Order".
            pattern = r"(?<![A-Za-z0-9])" + re.escape(key) + r"(?![A-Za-z0-9])"
            if re.search(pattern, text_lower):
                found[meta["canonical"]] = meta["category"]

        return [{"skill": skill, "category": category} for skill, category in sorted(found.items())]

    def extract_projects(self, projects_text: str) -> list:
        """Splits the projects section into individual project blocks."""
        if not projects_text:
            return []
        blocks = [b.strip() for b in re.split(r"\n{2,}", projects_text) if b.strip()]
        if len(blocks) <= 1:
            # Fall back to splitting on lines that don't start with '-' (likely titles)
            lines = [l for l in projects_text.split("\n") if l.strip()]
            projects, current = [], None
            for line in lines:
                if not line.strip().startswith("-"):
                    if current:
                        projects.append(current)
                    current = {"title": line.strip(), "description": []}
                elif current:
                    current["description"].append(line.lstrip("- ").strip())
            if current:
                projects.append(current)
            return projects

        projects = []
        for block in blocks:
            block_lines = block.split("\n")
            projects.append({
                "title": block_lines[0].strip(),
                "description": [l.lstrip("- ").strip() for l in block_lines[1:] if l.strip()],
            })
        return projects

    # ----------------------------------------------------------------
    # 4. Data structuring / orchestration
    # ----------------------------------------------------------------

    def parse_file(self, file_path: str) -> dict:
        """
        Main entry point. Ingests a resume file and returns a structured,
        JSON-serializable dictionary.
        """
        raw_text = self.extract_raw_text(file_path)
        clean = self.clean_text(raw_text)
        sections = self.segment_sections(clean)

        header_text = sections.get("header", "")
        contact_info = self.extract_contact_info(clean)
        name = self.extract_name(header_text if header_text else clean)

        result = {
            "file_name": os.path.basename(file_path),
            "parsed_at": datetime.utcnow().isoformat() + "Z",
            "candidate_name": name,
            "contact_info": contact_info,
            "summary": sections.get("summary"),
            "education": self.extract_education(sections.get("education", "")),
            "experience": self.extract_experience(sections.get("experience", "")),
            "projects": self.extract_projects(sections.get("projects", "")),
            "skills": self.extract_skills(clean),
            "certifications": sections.get("certifications"),
            "sections_detected": list(sections.keys()),
        }
        return result

    def parse_text(self, raw_text: str, file_name: str = "pasted_text") -> dict:
        """Same as parse_file but for already-extracted raw text (used by the UI)."""
        clean = self.clean_text(raw_text)
        sections = self.segment_sections(clean)
        header_text = sections.get("header", "")
        contact_info = self.extract_contact_info(clean)
        name = self.extract_name(header_text if header_text else clean)

        return {
            "file_name": file_name,
            "parsed_at": datetime.utcnow().isoformat() + "Z",
            "candidate_name": name,
            "contact_info": contact_info,
            "summary": sections.get("summary"),
            "education": self.extract_education(sections.get("education", "")),
            "experience": self.extract_experience(sections.get("experience", "")),
            "projects": self.extract_projects(sections.get("projects", "")),
            "skills": self.extract_skills(clean),
            "certifications": sections.get("certifications"),
            "sections_detected": list(sections.keys()),
        }


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python parser.py <path_to_resume.pdf|.docx>")
        sys.exit(1)

    rp = ResumeParser()
    parsed = rp.parse_file(sys.argv[1])
    print(json.dumps(parsed, indent=2))